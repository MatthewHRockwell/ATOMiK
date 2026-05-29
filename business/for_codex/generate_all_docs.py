#!/usr/bin/env python3
"""Generate all investor package documents in proper formats.
Run with: /home/mattrock/Projects/ATOMiK/.venv/bin/python generate_all_docs.py
"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).parent
REPO = ROOT.parent.parent

# ─── STYLE HELPERS ───────────────────────────────────────────────────────────

def styled_doc(title):
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    h = doc.add_heading(title, 0)
    h.runs[0].font.color.rgb = RGBColor(0x07, 0x0B, 0x12)
    h.runs[0].font.size = Pt(24)
    h.runs[0].bold = True

    sub = doc.add_paragraph("ATOMiK  ·  State-Aware Compute Architecture  ·  Pre-Seed  ·  May 2026")
    sub.runs[0].font.color.rgb = RGBColor(0x22, 0xD3, 0xEE)
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    return doc


def h1(doc, text):
    p = doc.add_heading(text, 1)
    p.runs[0].font.color.rgb = RGBColor(0x07, 0x0B, 0x12)
    return p


def h2(doc, text):
    p = doc.add_heading(text, 2)
    p.runs[0].font.color.rgb = RGBColor(0x1D, 0x32, 0x4A)
    return p


def body(doc, text, bold=False, size=11):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    p.runs[0].font.size = Pt(size)
    return p


def bullet(doc, text, color=None):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(10.5)
    if color:
        p.runs[0].font.color.rgb = color
    return p


def table_2col(doc, rows, header=None, widths=(2.8, 4.0)):
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    t.style = "Table Grid"
    if header:
        for i, h in enumerate(header):
            cell = t.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    offset = 1 if header else 0
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + offset].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    for col_idx, w in enumerate(widths):
        for row in t.rows:
            row.cells[col_idx].width = Inches(w)
    doc.add_paragraph()


def table_ncol(doc, rows, header=None, font_size=8.0):
    col_count = len(header) if header else len(rows[0])
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=col_count)
    t.style = "Table Grid"
    if header:
        for i, h in enumerate(header):
            cell = t.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    offset = 1 if header else 0
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + offset].cells[ci]
            cell.text = str(val)
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(font_size)
    doc.add_paragraph()


# ─── EXECUTIVE SUMMARY ───────────────────────────────────────────────────────

def gen_executive_summary():
    doc = styled_doc("Executive Summary")
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    h1(doc, "What ATOMiK Is")
    body(doc, "ATOMiK is a state-aware compute architecture that helps constrained edge, embedded, and AI-at-the-edge teams evaluate whether they can reduce wasted state movement by tracking meaningful change instead of repeatedly moving, scanning, syncing, replaying, or rebuilding full state.")
    body(doc, "The company is raising a $5M target pre-seed to move from hardware-backed primitive proof to measured customer workload proof, IP diligence, and ASIC feasibility. The round does not fund tape-out.", bold=True)

    h1(doc, "Why It Matters")
    bullet(doc, "Buyer pain: battery budget, enclosure heat, link budget, update latency, reliability, size, and weight.")
    bullet(doc, "First wedge: one state-heavy workload, one current baseline, and one painful constraint expensive enough to evaluate.")
    bullet(doc, "Why now: energy, bandwidth, thermal, and local-compute pressure are rising; IEA/LBNL figures are market context, not ATOMiK savings claims.")

    h1(doc, "Market Opportunity")
    table_2col(doc, [
        ("TAM context", "$1T+ 2026 semiconductor revenue backdrop. Context only; ATOMiK does not address all semiconductor spend."),
        ("SAM context", "$112B-$169B embedded systems market range from 2024 estimate to 2030 forecast; edge AI adds adjacent pressure."),
        ("Entry wedge", "$10M-$40M early annual revenue path from paid evaluations, design partners, and licensing if proof converts; not a ceiling or forecast."),
        ("IP/licensing context", "$8.14B 2025 to $11.2B 2029 semiconductor IP market context."),
    ], widths=(1.65, 5.15))

    h1(doc, "Evaluation Offer")
    table_2col(doc, [
        ("Give us", "One state-heavy workload, current baseline, and painful constraint."),
        ("We evaluate", "Where state movement creates waste and whether ATOMiK can improve that path."),
        ("You receive", "Workload map, baseline comparison, evidence map, fit/no-fit recommendation, and next-step plan."),
        ("Success", "Measured improvement against one agreed metric while preserving correctness."),
    ], widths=(1.55, 5.25))

    h1(doc, "Proof Today")
    table_2col(doc, [
        ("Zynq Desk v0.39-K", "HARDWARE_VALIDATED — live prototype UI proof image; not product maturity or performance proof."),
        ("Linux userspace to FPGA", "HARDWARE_VALIDATED — documented OS-to-bus path with 16/16 property checks."),
        ("AX7020 matrix", "LIVE_MEASURED — workload-specific wins and losses; quote with caveats."),
        ("Lean4-checked formal algebra", "FORMAL_PROOF where directly audited; otherwise SOFTWARE_VALIDATED / proof work present. Exact formal claims remain bounded to audited properties."),
        ("SD boot artifacts", "BUILD_ARTIFACT — build output exists; power-on run artifact still gated."),
    ], header=["Artifact", "Evidence Boundary"], widths=(2.1, 4.7))

    h1(doc, "Business Path")
    body(doc, "Paid proof reviews and technical evaluations -> design-partner evaluations -> IP/licensing diligence -> strategic licensing or platform partnership if workload proof supports it.")
    body(doc, "Defensibility wedge: Lean4-checked formal algebra, FPGA hardware validation, and an IP/proof registry. This is presented as a proof-bound diligence asset, not as a customer outcome claim.", size=9)
    body(doc, "Return narrative: pre-seed capital buys diligence assets a strategic partner can evaluate: workload proof, IP status, integration path, and ASIC feasibility. No exit, valuation, or return multiple is implied.", size=9)

    h1(doc, "Ask")
    table_2col(doc, [
        ("Target", "$5M pre-seed target; staged close terms counsel/CFO pending."),
        ("Use", "$1.5M engineering, $1.0M customer proof, $750K IP/legal, $750K ASIC feasibility, $1.0M finance/GTM/ops + reserve."),
        ("Need", "Capital, design-partner introductions, qualified workload access, and ASIC/IP diligence support."),
        ("Pending", "SAFE terms, valuation cap, discount, and close mechanics require CFO/counsel approval."),
    ], widths=(1.55, 5.25))

    h1(doc, "Evidence Boundary")
    body(doc, "Proof and financial claims remain artifact-bound or planning-only. Customer workload, production, revenue, and downstream battery, thermal, water, footprint, or ROI outcomes require matching measured artifacts or signed documents.", size=9)

    h1(doc, "Contact")
    body(doc, "Matthew Rockwell — Founder\nmatthew.h.rockwell@gmail.com", size=9)

    path = ROOT / "01_pitch_materials" / "ATOMiK_Executive_Summary.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── BUSINESS PLAN ───────────────────────────────────────────────────────────

def gen_business_plan():
    doc = styled_doc("ATOMiK Business Plan")
    body(doc, "CONTROLLED DILIGENCE - NOT FRONT-OF-ROOM - MAY 2026 - Founder-prepared; counsel/CFO review pending where applicable.", bold=True, size=9)

    h1(doc, "How To Read This Document")
    body(doc, "This plan is intended for controlled investor diligence after the initial pitch. It does not replace legal, financial, or technical diligence. Market figures are context only. Current proof claims are artifact-bound. Customer, revenue, battery, heat, water, cooling, and footprint outcomes require workload-specific measurement.", size=9.5)

    h1(doc, "Company Overview")
    body(doc, "ATOMiK is a state-aware compute architecture startup developing a hardware/IP primitive for reducing wasted state movement where workload evidence supports it. The company is pre-revenue, pre-seed stage, hardware-validated on FPGA for selected proof paths, and raising $5M to fund measured customer proof, IP diligence, ASIC feasibility, and commercial readiness.")
    body(doc, "Entity status: current operating entity is Rockwell Industries, LLC; intended VC-ready Delaware C-Corp structure remains counsel-pending before institutional close. See data room for current legal-status documentation.")

    h1(doc, "Problem")
    body(doc, "Many constrained systems - edge devices, embedded controllers, AI-at-the-edge systems, and remote industrial equipment - pay a hidden tax on state paths that repeatedly scan, compare, copy, move, sync, replay, or rebuild state even when only a compact change matters.")
    body(doc, "This waste manifests as:")
    bullet(doc, "Battery pressure in edge and IoT devices")
    bullet(doc, "Thermal pressure in sealed or fanless systems")
    bullet(doc, "Bandwidth pressure over intermittent or expensive links")
    bullet(doc, "Latency from rebuilding or replaying state that could have been tracked")
    bullet(doc, "Hardware overbuild to handle peak-state throughput that is mostly redundant")
    body(doc, "Market scale: Data-center energy 415->945 TWh by 2030 (IEA). U.S. data centers 176->325-580 TWh by 2028 (LBNL). 66B liters direct cooling water (2023 U.S.). Semiconductor sales were $791.7B in 2025 and Q1 2026 reached $298.5B (SIA).", size=9)
    body(doc, "Sources: IEA Energy and AI executive summary (https://www.iea.org/reports/energy-and-ai/executive-summary); LBNL 2024 United States Data Center Energy Usage Report (https://eta-publications.lbl.gov/sites/default/files/2024-12/lbnl-2024-united-states-data-center-energy-usage-report.pdf); SIA 2025 global sales release (https://www.semiconductors.org/global-annual-semiconductor-sales-increase-25-6-to-791-7-billion-in-2025/); SIA Q1 2026 sales release (https://www.semiconductors.org/global-semiconductor-sales-increase-25-from-q4-2025-to-q1-2026/). These figures are market context only, not ATOMiK savings claims.", size=7.8)

    h1(doc, "Solution: ATOMiK Architecture")
    body(doc, "Core equation: state = reference_state XOR accumulated_delta")
    body(doc, "Operations: LOAD (set reference state) → ACCUM (accumulate deltas via XOR) → READ (reconstruct current state) → SWAP (commit epoch boundary)")
    body(doc, "The primitive is self-inverse, commutative, associative, and identity-preserving under XOR. Formal proof work exists in the repo, and selected algebraic checks have passed through the Linux userspace-to-FPGA validation path.")
    body(doc, "What this means for customers: evaluate whether a constrained workload can move, scan, sync, or rebuild less state while preserving correctness.")

    h1(doc, "Current Proof Status")
    body(doc, "Hardware-validated proof exists for specific artifacts. The following is accurate as of May 2026:")
    bullet(doc, "ATOMiK Desk v0.39-K: framebuffer-native prototype UI running on live Zynq FPGA (HARDWARE_VALIDATED)")
    bullet(doc, "Algebraic tests: 16/16 PASS on XC7Z020 hardware through Linux userspace to FPGA (HARDWARE_VALIDATED)")
    bullet(doc, "AX7020 performance matrix: raw board-run artifacts with 4-way comparison (LIVE_MEASURED with documented caveats)")
    bullet(doc, "Formal proof foundation: FORMAL_PROOF where directly audited; otherwise proof work present. Avoid public proof counts until reconciled across repo, site, deck, and proof packet.")
    bullet(doc, "SD boot artifacts: BOOT.bin and bitstream exist; standalone power-on boot remains gated (BUILD_ARTIFACT)")
    body(doc, "IMPORTANT: Battery, heat, cooling, water, and footprint outcomes require workload-specific customer measurement. These are evaluation targets, not current claims.", bold=True, size=9)

    h1(doc, "Target Market & Customer Segments")
    table_2col(doc, [
        ("Edge / embedded", "Battery pressure, enclosure heat, intermittent links, reliability. Metrics: bytes moved, latency, bandwidth, power proxy."),
        ("AI at the edge", "Context/state movement, memory pressure, response time. Metrics: context retained, transfer avoided."),
        ("Remote/industrial/robotics/defense", "Weight, wattage, packet budget, field runtime. Metrics: runtime proxy, update cost."),
        ("Data center / infrastructure", "Power bill, cooling, water pressure, rack density. Metrics: bytes moved, power/thermal path. (Strategic expansion)"),
    ], header=["Segment", "Pain & Target Metrics"], widths=(2.0, 4.8))

    h1(doc, "Business Model")
    body(doc, "ATOMiK's near-term monetization is evaluation- and IP-led, not mass product manufacturing. Revenue path:")
    bullet(doc, "Paid technical evaluations — bring one workload, one baseline, one constraint; receive measured fit assessment")
    bullet(doc, "Design-partner engagements — deeper technical partnership with IP access")
    bullet(doc, "IP licensing — architecture and implementation licensed to chip/platform companies")
    bullet(doc, "Strategic optionality — licensing, integration partnership, or platform partnership if proof and IP diligence justify it")
    body(doc, "No revenue forecast until signed customer agreements. CFO approval required before projections are shared externally.")

    h1(doc, "Go-to-Market")
    body(doc, "First sale: one constrained team with one state-heavy workload hitting a hard limit. Entry point: technical evaluation. Proof-gate: measured improvement on one agreed metric while preserving correctness.")
    body(doc, "Not: outspend incumbents. Not: build a consumer product. Not: chase the data-center market before edge proof exists.")

    h1(doc, "Competitive Positioning")
    body(doc, "ATOMiK is not a CPU, GPU, NPU, or accelerator replacement. It is a state-aware architecture layer for workloads where state movement is the binding constraint.")
    table_2col(doc, [
        ("General CPUs", "Flexible execution. ATOMiK differs by evaluating whether the workload can be represented as reference state plus accumulated change."),
        ("Compression", "Reduces representation size. ATOMiK targets when state movement or repeated reconstruction can be avoided."),
        ("Caching", "Keeps hot data nearby. ATOMiK targets meaningful change and epoch boundaries."),
        ("Custom accelerators", "Optimize known kernels. ATOMiK is an architecture primitive that still requires workload-specific proof."),
    ], header=["Comparison", "Differentiation"], widths=(1.8, 5.0))

    h1(doc, "IP & Legal")
    body(doc, "Provisional patent filed (see data room: 03_intellectual_property). Patent conversion, prior-art review, trade-secret controls, and final IP assignment chain remain counsel-review workstreams. Pre-seed allocates $750K to IP/legal.")

    h1(doc, "Team")
    body(doc, "Matthew Rockwell — Founder. Hardware engineer, FPGA/ASIC development background. Built the ATOMiK architecture, validated on physical hardware. See data room for full founder profile.")
    body(doc, "Fractional CFO: needed before SAFE terms finalized. First technical hire and advisory board in progress.")

    h1(doc, "Financial Summary")
    body(doc, "See ATOMiK_Financial_Model.xlsx for full detail. Summary:")
    table_2col(doc, [
        ("Pre-seed target", "$5M target; final staged close and SAFE mechanics CFO/counsel pending"),
        ("Instrument", "Planned SAFE — final terms require CFO/counsel"),
        ("Runway", "18-24 month target plan at $5M raise"),
        ("Engineering", "$1.5M (demo hardening, measurement tooling)"),
        ("Customer proof", "$1.0M (evaluation SOWs, baseline harnesses)"),
        ("IP / legal", "$750K (patent conversion, diligence packet)"),
        ("ASIC feasibility", "$750K (expert review, go/no-go)"),
        ("Finance / GTM", "$600K + $400K reserve"),
        ("Tape-out", "NOT funded in this round"),
        ("Revenue", "No forecast without signed agreements"),
    ], widths=(2.2, 4.6))

    h1(doc, "Risks & Mitigations")
    table_2col(doc, [
        ("Customer validation pending", "Budget dedicated to proof reviews, technical evaluations, and design-partner outreach."),
        ("Battery/thermal/water outcomes unproven", "Evaluation targets. Require workload measurement. Not claimed yet."),
        ("ASIC economics unvalidated", "Funded in budget. Expert review before tape-out commitment."),
        ("IP protection coverage", "Funded conversion. Prior-art review in progress."),
        ("Incumbent response", "Build IP and proof quickly. Pursue partner conversations without implying a promised exit."),
    ], header=["Risk", "Mitigation"], widths=(2.4, 4.4))

    path = ROOT / "01_pitch_materials" / "ATOMiK_Business_Plan.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── ONE PAGER ───────────────────────────────────────────────────────────────

def gen_one_pager():
    doc = styled_doc("ATOMiK — One Page Overview")
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    body(doc, "ATOMiK makes change the unit of compute.", bold=True, size=13)
    body(doc, "State-aware compute evaluation for edge and embedded teams constrained by battery, heat, bandwidth, latency, reliability, size, weight, or hardware footprint.", size=8.8)

    h2(doc, "Problem")
    body(doc, "Customers already feel the pain: batteries die too quickly, hardware is maxed out, systems run hot, devices get too big or heavy, costs rise, and more performance is needed. Before they buy bigger hardware, batteries, cooling, bandwidth, or redesign, ATOMiK tests whether a constrained state path can do less wasted work.", size=8.5)

    h2(doc, "Evaluation Offer")
    body(doc, "Bring one state-heavy workload, one current baseline, and one painful constraint. ATOMiK returns a workload map, baseline comparison, evidence map, fit/no-fit recommendation, and next-step plan. If there is no fit, we say so.", size=8.5)

    h2(doc, "Market And Wedge")
    body(doc, "$1T+ semiconductor backdrop; $112B-$169B embedded systems SAM context; $10M-$40M entry wedge from paid evaluations, design partners, and licensing if proof converts. This is an early wedge, not a ceiling or forecast.", size=8.4)

    h2(doc, "Proof And Defensibility")
    body(doc, "Proof today: Zynq Desk v0.39-K HARDWARE_VALIDATED UI artifact; Linux userspace-to-FPGA HARDWARE_VALIDATED path with 16/16 checks; AX7020 LIVE_MEASURED matrix with wins/losses/caveats; Lean4-checked formal algebra with FORMAL_PROOF only where audited. Defensibility wedge: formal algebra + FPGA hardware validation + IP/proof registry.", size=8.2)

    h2(doc, "Pre-Seed Ask")
    body(doc, "$5M target pre-seed; final SAFE mechanics counsel/CFO pending. Use of funds: $1.5M engineering, $1.0M customer proof, $750K IP/legal, $750K ASIC feasibility, $1.0M finance/GTM/ops + reserve. Milestone: measured workload proof + IP packet + ASIC feasibility; no tape-out.", size=8.3)

    body(doc, "Evidence boundary: proof and financial claims remain artifact-bound or planning-only. Customer workload, production, revenue, and downstream battery, thermal, water, footprint, or ROI outcomes require matching measured artifacts or signed documents.", size=7.7)
    body(doc, "matthew.h.rockwell@gmail.com  |  May 2026", size=8.0)

    path = ROOT / "01_pitch_materials" / "ATOMiK_One_Pager.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── FINANCIAL MODEL (EXCEL) ─────────────────────────────────────────────────

def gen_financial_model():
    wb = openpyxl.Workbook()

    DARK = "070B12"; CYAN = "22D3EE"; GREEN = "22C55E"; AMBER = "F59E0B"; MUTED = "9FB1C7"
    LIGHT = "F4F8FF"; PANEL = "0D1420"; ROW_A = "111827"; ROW_B = "0D1420"; RED = "F87171"
    MONEY = '$#,##0;($#,##0);-'; MONEY_DEC = '$#,##0.00;($#,##0.00);-'; PCT = '0.0%'; NUM = '#,##0;-#,##0;-'

    thin = Side(style="thin", color="233044")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def setup(ws, widths):
        for col, width in widths.items():
            ws.column_dimensions[col].width = width
        ws.freeze_panes = "A3"
        ws.sheet_view.showGridLines = False

    def hdr(ws, row, col, text, bold=True, bg=DARK, fg=CYAN, sz=11):
        c = ws.cell(row, col, text)
        c.font = Font(name="Calibri", bold=bold, color=fg, size=sz)
        c.fill = PatternFill("solid", fgColor=bg)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = border
        return c

    def val(ws, row, col, v, fmt=None, bold=False, fg=LIGHT, bg=None, align=None):
        c = ws.cell(row, col, v)
        c.font = Font(name="Calibri", bold=bold, color=fg, size=10)
        if bg: c.fill = PatternFill("solid", fgColor=bg)
        if fmt: c.number_format = fmt
        if align is None:
            align = "right" if isinstance(v, (int, float)) or (isinstance(v, str) and v.startswith("=")) else "left"
        c.alignment = Alignment(horizontal=align, vertical="top", wrap_text=True)
        c.border = border
        return c

    def title(ws, text, last_col):
        hdr(ws, 1, 1, text, fg=CYAN, bg=DARK, sz=14)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=last_col)

    def section_note(ws, row, text, last_col, fg=AMBER):
        c = val(ws, row, 1, text, bold=True, fg=fg, bg=DARK, align="left")
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=last_col)
        c.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        return c

    scenarios = {
        "Staged close": {"amount": 3_000_000, "reserve": 250_000, "months": 18, "note": "Staged path: proof, IP conversion, demo hardening, first measured evaluations."},
        "Target": {"amount": 5_000_000, "reserve": 400_000, "months": 24, "note": "Aggie-facing live ask: workload proof, IP packet, ASIC feasibility, customer validation, and operating capacity."},
        "Expanded": {"amount": 7_500_000, "reserve": 600_000, "months": 24, "note": "Expanded validation and hiring support; still not tape-out funding."},
    }
    categories = [
        ("Engineering & Demo", 900_000, 1_500_000, 2_250_000, "Repeatable investor/design-partner proof system."),
        ("Customer Proof", 600_000, 1_000_000, 1_500_000, "Measured workload artifacts with correctness preserved."),
        ("IP / Legal", 450_000, 750_000, 1_125_000, "Counsel-reviewed IP packet and patent-conversion plan."),
        ("ASIC Feasibility", 450_000, 750_000, 1_125_000, "Go/no-go feasibility report; quote-backed path; no tape-out funding."),
        ("Finance / GTM / Ops", 350_000, 600_000, 900_000, "CFO support, investor reporting, design-partner outreach, entity/admin readiness."),
        ("Reserve", 250_000, 400_000, 600_000, "Runway buffer; do not convert expanded plan into overbuild."),
    ]

    # Sheet 1: README / Assumptions / Model Map
    ws_intro = wb.active; ws_intro.title = "README Model Map"
    setup(ws_intro, {'A':28,'B':72})
    title(ws_intro, "README / Assumptions / Model Map", 2)
    model_map_rows = [
        ("Model version/date", "2026-05-27 financial diligence draft."),
        ("Owner", "Founder-prepared; pending CFO/counsel review."),
        ("Intended use", "Controlled diligence and investor follow-up. Not a public/front-of-room handout by default."),
        ("Not advice", "Not legal, tax, accounting, securities, or valuation advice."),
        ("Core raise assumption", "$5.0M target pre-seed; $3.0M staged close scenario; $7.5M expanded scenario."),
        ("Runway assumption", "Target plan models roughly 24 months: $4.60M planned spend plus $400K reserve."),
        ("Revenue assumption", "No revenue forecast. Reservation pricing is a qualification/scoping signal only."),
        ("Investor-facing sheets", "Use of Funds; Runway & Burn; Dilution Sensitivity; Milestone Gates; Market Context; Caveats."),
        ("Internal/control sheets", "18-Month Cash Plan; Evaluation Pricing; Evaluation SOW Economics; Cap Table Draft; Unmodeled / CFO Needs."),
        ("Do not change without review", "Raise scenarios, SAFE sensitivity, source-backed market figures, and claim guardrails."),
        ("Source register", "See ATOMiK_Financial_Source_Register and Market Context source URL column."),
        ("Open diligence gaps", "CFO/counsel terms, option pool, taxes/benefits, insurance, legal close costs, vendor quotes, and first signed SOW."),
    ]
    for i,h in enumerate(["Item","Current treatment"],1): hdr(ws_intro,3,i,h,fg=LIGHT)
    for r,row in enumerate(model_map_rows,4):
        bg = ROW_A if r % 2 else ROW_B
        val(ws_intro,r,1,row[0],bold=True,bg=bg)
        val(ws_intro,r,2,row[1],bg=bg)
    section_note(ws_intro, 18, "Use this workbook as a controlled milestone model. Do not distribute as final CFO-approved terms.", 2)

    # Sheet 2: Use of Funds
    ws = wb.create_sheet("Use of Funds")
    setup(ws, {'A':28,'B':15,'C':15,'D':15,'E':12,'F':50})
    title(ws, "ATOMiK - Pre-Seed Use of Funds (Formula-Backed)", 6)
    headers = ["Category", "Staged $3.0M", "Target $5.0M", "Expanded $7.5M", "Target %", "Proof Gate / Purpose"]
    for i, h in enumerate(headers, 1): hdr(ws, 2, i, h, fg=LIGHT)
    for i, (name, mn, tgt, sx, purpose) in enumerate(categories, 3):
        bg = ROW_A if i % 2 else ROW_B
        val(ws, i, 1, name, bold=True, bg=bg)
        val(ws, i, 2, mn, MONEY, fg=GREEN, bg=bg)
        val(ws, i, 3, tgt, MONEY, fg=GREEN, bg=bg)
        val(ws, i, 4, sx, MONEY, fg=GREEN, bg=bg)
        val(ws, i, 5, f"=C{i}/$C$9", PCT, fg=CYAN, bg=bg)
        val(ws, i, 6, purpose, bg=bg)
    total_row = 9
    val(ws, total_row, 1, "TOTAL", bold=True, fg=CYAN, bg=DARK)
    for col in [2,3,4]:
        val(ws, total_row, col, f"=SUM({get_column_letter(col)}3:{get_column_letter(col)}8)", MONEY, bold=True, fg=CYAN, bg=DARK)
    val(ws, total_row, 5, "=SUM(E3:E8)", PCT, bold=True, fg=CYAN, bg=DARK)
    val(ws, total_row, 6, "Funds proof, IP, feasibility, operations. Does not fund tape-out.", bold=True, fg=AMBER, bg=DARK)
    section_note(ws, 11, "CFO diligence note: amounts are milestone-planning allocations. Vendor quotes, hiring plan, tax/accounting treatment, and legal close mechanics still require approval.", 6)

    # Sheet 2: Runway & Burn
    ws2 = wb.create_sheet("Runway & Burn")
    setup(ws2, {'A':18,'B':14,'C':14,'D':16,'E':14,'F':18,'G':22,'H':48})
    title(ws2, "Runway And Burn Sensitivity", 8)
    for i,h in enumerate(["Scenario","Raise","Reserve","Spendable Cash","Planning Months","Gross $/Month","Ex-Reserve $/Month","Notes"],1): hdr(ws2,2,i,h,fg=LIGHT)
    for r, (scenario, data) in enumerate(scenarios.items(), 3):
        bg = ROW_A if r % 2 else ROW_B
        val(ws2,r,1,scenario,bold=True,bg=bg)
        val(ws2,r,2,data["amount"],MONEY,fg=GREEN,bg=bg)
        val(ws2,r,3,data["reserve"],MONEY,fg=AMBER,bg=bg)
        val(ws2,r,4,f"=B{r}-C{r}",MONEY,fg=GREEN,bg=bg)
        val(ws2,r,5,data["months"],NUM,fg=CYAN,bg=bg)
        val(ws2,r,6,f"=B{r}/E{r}",MONEY,fg=GREEN,bg=bg)
        val(ws2,r,7,f"=D{r}/E{r}",MONEY,fg=GREEN,bg=bg)
        val(ws2,r,8,data["note"],bg=bg)
    section_note(ws2,7,"Target-plan monthly budget by category. These are average equivalents, not a monthly hiring commitment; spend should be staged by milestone.",8)
    for i,h in enumerate(["Category","Target Allocation","Months","Avg $/Month","Gate"],1): hdr(ws2,8,i,h,fg=LIGHT)
    for idx, (name, _mn, tgt, _sx, gate) in enumerate(categories, 9):
        bg = ROW_A if idx % 2 else ROW_B
        val(ws2,idx,1,name,bold=True,bg=bg)
        val(ws2,idx,2,tgt,MONEY,fg=GREEN,bg=bg)
        val(ws2,idx,3,"='Runway & Burn'!$E$4",NUM,fg=CYAN,bg=bg)
        val(ws2,idx,4,f"=B{idx}/C{idx}",MONEY,fg=GREEN,bg=bg)
        val(ws2,idx,5,gate,bg=bg)
    val(ws2,16,1,"Formula check: target allocation total",bold=True,fg=CYAN,bg=DARK)
    val(ws2,16,2,"=SUM(B9:B14)",MONEY,bold=True,fg=CYAN,bg=DARK)
    val(ws2,16,4,"=SUM(D9:D14)",MONEY,bold=True,fg=CYAN,bg=DARK)
    val(ws2,16,5,"Monthly total should equal target gross/month.",fg=AMBER,bg=DARK)
    ws2.merge_cells('E16:H16')

    # Sheet 3: Monthly Cash Plan
    ws_month = wb.create_sheet("18-Month Cash Plan")
    month_widths = {'A':32}
    for idx in range(2, 22):
        month_widths[get_column_letter(idx)] = 12
    setup(ws_month, month_widths)
    title(ws_month, "18-Month Target Cash Plan - Planning Draft", 21)
    section_note(ws_month, 2, "Target-plan monthly budget. Values are planning allocations, not approved hiring commitments. Spend should be staged behind proof gates and vendor quotes.", 21)
    months = [f"M{i}" for i in range(1, 19)]
    hdr(ws_month, 4, 1, "Line Item", fg=LIGHT)
    for c, m in enumerate(months, 2):
        hdr(ws_month, 4, c, m, fg=LIGHT)
    hdr(ws_month, 4, 20, "Total", fg=LIGHT)
    hdr(ws_month, 4, 21, "Notes", fg=LIGHT)
    monthly_rows = [
        ("Engineering capacity / contractors", [20_000, 20_000] + [30_000] * 16, "FPGA/ASIC and workload-validation capacity; hire timing staged by close."),
        ("Demo hardware, tools, packaging", [20_000, 20_000, 10_000, 10_000, 10_000, 10_000] + [0] * 12, "Boards, test equipment, software/tooling, investor-demo reliability."),
        ("Customer workload proof", [5_000, 5_000] + [20_000] * 4 + [30_000] * 6 + [21_666.6667] * 6, "Evaluation harnesses, traces, baseline work, proof artifacts."),
        ("IP / legal", [30_000] * 4 + [25_000] * 4 + [20_000] * 4 + [0] * 6, "Patent conversion, prior art, entity/IP cleanup, diligence packet."),
        ("ASIC feasibility", [5_000] * 3 + [10_000] * 3 + [30_000] * 6 + [12_500] * 6, "Mentor, feasibility scoping, external review; no tape-out."),
        ("Finance / GTM / ops", [250_000 / 18] * 18, "Fractional CFO/accounting, investor reporting, outreach, insurance/admin."),
    ]
    for r, (name, values, note) in enumerate(monthly_rows, 5):
        bg = ROW_A if r % 2 else ROW_B
        val(ws_month, r, 1, name, bold=True, bg=bg)
        for c, amount in enumerate(values, 2):
            val(ws_month, r, c, amount, MONEY, fg=GREEN, bg=bg)
        val(ws_month, r, 20, f"=SUM(B{r}:S{r})", MONEY, bold=True, fg=CYAN, bg=bg)
        val(ws_month, r, 21, note, bg=bg)
    total_r = 12
    val(ws_month, total_r, 1, "Monthly planned spend", bold=True, fg=CYAN, bg=DARK)
    for c in range(2, 20):
        val(ws_month, total_r, c, f"=SUM({get_column_letter(c)}5:{get_column_letter(c)}10)", MONEY, bold=True, fg=CYAN, bg=DARK)
    val(ws_month, total_r, 20, "=SUM(B12:S12)", MONEY, bold=True, fg=CYAN, bg=DARK)
    val(ws_month, total_r, 21, "Should equal $4.60M planned spend, excluding $400K reserve.", fg=AMBER, bg=DARK)
    cum_r = 13
    val(ws_month, cum_r, 1, "Cumulative spend", bold=True, fg=CYAN, bg=DARK)
    for c in range(2, 20):
        col = get_column_letter(c)
        val(ws_month, cum_r, c, f"=SUM($B$12:{col}$12)", MONEY, bold=True, fg=CYAN, bg=DARK)
    val(ws_month, cum_r, 20, "=S13", MONEY, bold=True, fg=CYAN, bg=DARK)
    val(ws_month, cum_r, 21, "Tracks spend against target spendable cash.", fg=AMBER, bg=DARK)
    rem_r = 14
    val(ws_month, rem_r, 1, "Remaining spendable cash", bold=True, fg=GREEN, bg=DARK)
    for c in range(2, 20):
        col = get_column_letter(c)
        val(ws_month, rem_r, c, f"=IF(ABS(1850000-{col}13)<1,0,ROUND(1850000-{col}13,0))", MONEY, bold=True, fg=GREEN, bg=DARK)
    val(ws_month, rem_r, 20, "=S14", MONEY, bold=True, fg=GREEN, bg=DARK)
    val(ws_month, rem_r, 21, "Target raise minus reserve is $1.85M.", fg=AMBER, bg=DARK)
    res_r = 15
    val(ws_month, res_r, 1, "Reserve held", bold=True, fg=AMBER, bg=DARK)
    for c in range(2, 20):
        val(ws_month, res_r, c, 150_000, MONEY, bold=True, fg=AMBER, bg=DARK)
    val(ws_month, res_r, 20, 150_000, MONEY, bold=True, fg=AMBER, bg=DARK)
    val(ws_month, res_r, 21, "Held as buffer; not included in planned spend.", fg=AMBER, bg=DARK)
    section_note(ws_month, 17, "Cash-plan check: B12:S12 totals should equal the active target spend and row 15 should preserve the current reserve. Add taxes, benefits, and exact vendor quotes once CFO/accounting has current assumptions.", 21)

    # Sheet 4: Evaluation Pricing
    ws3 = wb.create_sheet("Evaluation Pricing")
    setup(ws3, {'A':30,'B':18,'C':40,'D':34,'E':40})
    title(ws3, "Evaluation Pricing And Revenue Signals - Planning Only", 5)
    section_note(ws3,2,"This sheet is not a revenue forecast. Reservation prices qualify interest and may be credited toward larger scoped work; larger evaluations require written SOWs.",5)
    for i,h in enumerate(["Offer","Public / Planning Price","External Wording","Financial Treatment","Caveat"],1): hdr(ws3,4,i,h,fg=LIGHT)
    pricing_rows = [
        ("Proof Review Reservation", 750, "One-time reservation for proof/workload review.", "Small qualification payment; not material revenue.", "Does not create commercial license or prove customer demand by itself."),
        ("Technical Evaluation Reservation", 2500, "One-time reservation for evaluation scoping.", "Small qualification payment; may convert into scoped SOW.", "Final deliverables and terms confirmed separately."),
        ("Scoped Design Partner Evaluation", "Request-based", "Written scope around workload, baseline, metric, deliverables.", "Potential first material customer-funded proof path.", "No forecast until signed SOW."),
        ("Licensing / IP Diligence", "Request-based", "Partner-specific diligence path.", "Future/mid-term revenue candidate.", "Requires workload proof, IP packet, and integration path."),
        ("Commercial Licensing / Embedded IP", "Future", "Not a pre-seed revenue forecast.", "Strategic upside path.", "Do not quote license pricing before customer/partner evidence."),
    ]
    for r,row in enumerate(pricing_rows,5):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1):
            fmt = MONEY if isinstance(v,(int,float)) else None
            fg = GREEN if isinstance(v,(int,float)) else LIGHT
            val(ws3,r,c,v,fmt,fg=fg,bg=bg)
    section_note(ws3,12,"Illustrative reservation cash only - not a forecast and not a valuation input.",5)
    for i,h in enumerate(["Input","Conservative","Base","Upside","Notes"],1): hdr(ws3,13,i,h,fg=LIGHT)
    assumptions = [
        ("Proof review reservations", 0, 2, 5, "Uses current public $750 reservation."),
        ("Technical eval reservations", 0, 1, 3, "Uses current public $2,500 reservation."),
        ("Reservation cash", "=B14*750+B15*2500", "=C14*750+C15*2500", "=D14*750+D15*2500", "Qualification signal only; serious revenue requires SOW."),
        ("Scoped SOW revenue", "TBD", "TBD", "TBD", "Do not forecast until a written scope exists."),
    ]
    for r,row in enumerate(assumptions,14):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1):
            fmt = MONEY if (isinstance(v,str) and v.startswith("=")) else None
            val(ws3,r,c,v,fmt,fg=GREEN if fmt else LIGHT,bg=bg)

    # Sheet 5: Evaluation SOW Economics
    ws_sow = wb.create_sheet("Evaluation SOW Economics")
    setup(ws_sow, {'A':30,'B':36,'C':34,'D':42})
    title(ws_sow, "Evaluation SOW Economics - Planning Only", 4)
    section_note(ws_sow, 2, "Planning template only. Do not forecast revenue until a signed SOW defines price, scope, deliverables, ownership, and acceptance criteria.", 4)
    for i,h in enumerate(["Workstream","What it covers","Commercial treatment","Diligence caveat"],1): hdr(ws_sow,4,i,h,fg=LIGHT)
    sow_rows = [
        ("Workload intake", "Representative workload, baseline, painful constraint, and target environment.", "TBD / scoped quote only.", "Not forecast until customer supplies scope and signs terms."),
        ("Baseline mapping", "Current implementation, state path, trace/logs, and agreed measurement method.", "May be included in scoped evaluation fee.", "Requires customer artifacts; no customer-value claim without baseline."),
        ("Success metric definition", "One primary metric and decision threshold.", "Included in evaluation scope.", "Metric must preserve correctness and map to buyer pain."),
        ("Test harness / artifacts", "Trace replay, baseline vs ATOMiK path, raw logs, proof card, and claims entry.", "Potential material customer-funded proof path.", "No public claim unless correctness passes and artifacts are reproducible."),
        ("Readout", "Workload map, baseline comparison, evidence map, fit/no-fit recommendation.", "Deliverable of technical evaluation.", "Fit/no-fit is a valid outcome; do not promise positive results."),
        ("Partner-funded proof work", "Optional deeper integration after proof review.", "Request-based SOW.", "Requires counsel review for IP, confidentiality, and ownership."),
        ("Accounting treatment", "Revenue recognition and booking treatment.", "Pending CFO/accounting.", "Do not model until signed SOW and accounting review."),
    ]
    for r,row in enumerate(sow_rows,5):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1): val(ws_sow,r,c,v,bg=bg)

    # Sheet 6: Milestone Gates
    ws4 = wb.create_sheet("Milestone Gates")
    setup(ws4, {'A':18,'B':36,'C':28,'D':28,'E':42})
    title(ws4, "18-Month Milestone Gates", 5)
    for i,h in enumerate(["Window","Milestone","Budget Owner","Evidence Gate","Investor Readout"],1): hdr(ws4,2,i,h,fg=LIGHT)
    gates = [
        ("0-30 days", "Lock financing structure, counsel plan, design-partner target list.", "Finance / Legal", "CFO/counsel reviewed plan.", "Terms are ready to circulate; no improvised cap."),
        ("30-60 days", "Harden Zynq demo into repeatable package.", "Engineering", "Demo checklist and current artifact bundle.", "Prototype proof is lower-friction and repeatable."),
        ("60-90 days", "Define first customer workload and baseline plan.", "Customer Proof", "Evaluation SOW and one success metric.", "Qualified workload exists or no-fit is documented."),
        ("90-180 days", "Produce first measured workload artifact.", "Customer Proof / Engineering", "Correctness-preserving workload result.", "Evidence-backed customer-value claim for one workload."),
        ("6-9 months", "Convert artifact into paid evaluation, LOI, or no-fit decision.", "GTM / Customer Proof", "Customer decision record.", "Early commercial signal or disciplined learning."),
        ("9-12 months", "Complete ASIC feasibility scope and IP diligence packet.", "ASIC / Legal", "Mentor-reviewed feasibility and counsel-reviewed IP memo.", "Silicon path is scoped without tape-out spend."),
        ("12-18 months", "Prepare next financing or strategic partner process.", "Finance / GTM", "Measured proof, partner pipeline, IP packet.", "Next-round or strategic diligence package."),
    ]
    for r,row in enumerate(gates,3):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1): val(ws4,r,c,v,bg=bg)

    # Sheet 7: Dilution Sensitivity
    ws5 = wb.create_sheet("Dilution Sensitivity")
    setup(ws5, {'A':18,'B':18,'C':18,'D':18,'E':22,'F':46})
    title(ws5, "Post-Money SAFE Dilution Sensitivity - Illustrative Only", 6)
    section_note(ws5,2,"Not a valuation recommendation. Shows ownership sold under illustrative SAFE cap math before option-pool changes, prior instruments, side letters, or priced-round mechanics.",6)
    for i,h in enumerate(["Post-Money Cap","$3.0M Staged","$5.0M Target","$7.5M Expanded","Founder Before Pool (Target)","Use In Room"],1): hdr(ws5,4,i,h,fg=LIGHT)
    caps = [8_000_000, 10_000_000, 12_000_000, 15_000_000, 18_000_000]
    for r,cap in enumerate(caps,5):
        bg = ROW_A if r % 2 else ROW_B
        val(ws5,r,1,cap,MONEY,fg=GREEN,bg=bg)
        val(ws5,r,2,f"=1250000/A{r}",PCT,fg=CYAN,bg=bg)
        val(ws5,r,3,f"=2000000/A{r}",PCT,fg=CYAN,bg=bg)
        val(ws5,r,4,f"=2750000/A{r}",PCT,fg=CYAN,bg=bg)
        val(ws5,r,5,f"=1-C{r}",PCT,fg=CYAN,bg=bg)
        val(ws5,r,6,"CFO/counsel must set actual cap; do not offer in room.",fg=AMBER,bg=bg)
    section_note(ws5,12,"Rule of thumb: at a $25M post-money cap, a $5.0M SAFE implies 20.0% SAFE ownership before option-pool and later-round effects. This is why cap discipline matters.",6)

    # Sheet 8: Market Context
    ws6 = wb.create_sheet("Market Context")
    setup(ws6, {'A':28,'B':24,'C':26,'D':34,'E':56})
    title(ws6, "Market Context - Source-Backed Reference Data", 5)
    for i,h in enumerate(["Metric","Current","Projected / Context","Source / Use","Primary source URL"],1): hdr(ws6,2,i,h,fg=LIGHT)
    market_data = [
        ("Global data-center energy", "415 TWh (2024)", "~945 TWh by 2030", "IEA Energy and AI. Why-now context only.", "https://www.iea.org/reports/energy-and-ai/energy-demand-from-ai"),
        ("U.S. data-center energy", "176 TWh (2023)", "325-580 TWh (2028)", "LBNL 2024 U.S. Data Center Report.", "https://eta-publications.lbl.gov/sites/default/files/2024-12/lbnl-2024-united-states-data-center-energy-usage-report_1.pdf"),
        ("U.S. data-center cooling water", "66B liters (2023)", "60-124B liters hyperscale direct use (2028)", "LBNL. Water savings require site-specific proof.", "https://eta-publications.lbl.gov/sites/default/files/2024-12/lbnl-2024-united-states-data-center-energy-usage-report_1.pdf"),
        ("Global semiconductor sales", "$791.7B (2025)", "Q1 2026 sales $298.5B", "SIA. Backdrop only; not ATOMiK TAM.", "https://www.semiconductors.org/global-annual-semiconductor-sales-increase-25-6-to-791-7-billion-in-2025/; https://www.semiconductors.org/global-semiconductor-sales-increase-25-from-q4-2025-to-q1-2026/"),
        ("U.S. pre-seed market", "~3,000 Carta startups / >$2.3B in Q1 2026", "Expected final total around $2.9B", "Carta. Financing context only.", "https://carta.com/sg/en/data/state-of-pre-seed-q1-2026/"),
        ("U.S. seed valuation benchmark", "$18.4M median pre-money", "Q1 2026", "PitchBook/NVCA. Not ATOMiK valuation.", "https://nvca.org/wp-content/uploads/2026/04/Q1-2026-PitchBook-NVCA-Venture-Monitor.pdf"),
        ("U.S. seed deal-size benchmark", "$3.0M median deal value", "Q1 2026", "PitchBook/NVCA. Supports disciplined staged financing context.", "https://nvca.org/wp-content/uploads/2026/04/Q1-2026-PitchBook-NVCA-Venture-Monitor.pdf"),
        ("Accelerator comparables", "YC $500K; Techstars $220K", "Different program structures", "Dilution context; not direct round comps.", "https://www.ycombinator.com/deal; https://www.techstars.com/newsroom/investment-terms"),
        ("ASIC cost discipline", "NRE varies by design, node, IP, EDA, mask, packaging, test", "Use feasibility-first framing", "AnySilicon. Avoid false precision before expert quote.", "https://anysilicon.com/semiconductor-manufacturing-cost-breakdown/; https://anysilicon.com/asic-nre-explained/"),
    ]
    for r,row in enumerate(market_data,3):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1): val(ws6,r,c,v,bg=bg)
    section_note(ws6,13,"Market figures support why-now and financing context. They do not prove ATOMiK savings, customer demand, or valuation.",5)

    # Sheet 9: Cap Table Draft
    ws7 = wb.create_sheet("Cap Table Draft")
    setup(ws7, {'A':30,'B':20,'C':18,'D':22,'E':54})
    title(ws7, "Cap Table - Draft Placeholder / Counsel Review Required", 5)
    for i,h in enumerate(["Holder","Security","Shares","Current Status","Required Before Close"],1): hdr(ws7,3,i,h,fg=LIGHT)
    cap_rows = [
        ("Matthew Rockwell (Founder)", "Common", "TBD", "Founder / pre-close", "Incorporation, stock issuance, founder IP assignment."),
        ("ESOP / Option Pool", "Common reserved", "TBD", "To be created", "CFO/counsel sizing; likely 10-15% discussion, not commitment."),
        ("Pre-Seed Investors", "Planned SAFE", "N/A", "Planned", "Final SAFE terms, cap, discount, pro-rata, side letters."),
        ("Prior instruments", "TBD", "TBD", "Must verify", "Confirm no hidden notes, grants, IP obligations, or advisor promises."),
    ]
    for r,row in enumerate(cap_rows,4):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1): val(ws7,r,c,v,bg=bg)
    section_note(ws7,10,"Do not share this as a formal cap table. It is a diligence checklist until entity formation, stock issuance, and counsel-reviewed records are complete.",5)

    # Sheet 10: Unmodeled CFO Needs
    ws_cfo = wb.create_sheet("Unmodeled CFO Needs")
    setup(ws_cfo, {'A':34,'B':44,'C':42})
    title(ws_cfo, "Unmodeled / CFO Estimate Needed", 3)
    section_note(ws_cfo, 2, "These are visible diligence gaps, not hidden costs. Estimate with CFO/accounting/counsel before final model distribution.", 3)
    for i,h in enumerate(["Item","Why it matters","Current treatment"],1): hdr(ws_cfo,4,i,h,fg=LIGHT)
    cfo_rows = [
        ("Payroll taxes", "Affects fully loaded burn.", "Not separately modeled; included only as broad finance/ops planning."),
        ("Benefits", "Hiring costs differ materially by role and timing.", "Requires CFO/accounting estimate before hiring plan is shared."),
        ("Insurance", "D&O, general liability, cyber, and hardware/customer evaluation coverage may be needed.", "Needs broker quote."),
        ("Accounting/bookkeeping/tax", "Ongoing reporting and investor-readiness cost.", "Broadly included in Finance/GTM/Ops; estimate needed."),
        ("Delaware C-Corp conversion", "Required for VC-ready close path if counsel chooses conversion/formation.", "Legal budget allocation exists; exact cost pending counsel."),
        ("Legal close costs", "SAFE review, side letters, cap-table setup, and closing mechanics.", "Pending counsel/CFO."),
        ("IP prosecution variability", "Patent conversion and prior-art work can vary.", "IP/legal allocation exists; quote needed."),
        ("FPGA/dev-board/tooling", "Hardware validation may need boards, instruments, tools, and licenses.", "Demo/tooling allocation exists; vendor quotes needed."),
        ("Cloud/dev infrastructure", "Build, storage, collaboration, and artifact hosting.", "Not separately forecast."),
        ("Travel/customer evaluation support", "Design-partner work may require on-site support.", "Not separately forecast."),
        ("Additional contingency", "Reserve may need adjustment after quotes.", "$400K reserve held; not CFO-approved contingency model."),
    ]
    for r,row in enumerate(cfo_rows,5):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1): val(ws_cfo,r,c,v,bg=bg)

    # Sheet 11: Caveats
    ws8 = wb.create_sheet("Caveats")
    setup(ws8, {'A':34,'B':74})
    title(ws8, "Financial Diligence Caveats", 2)
    caveats = [
        ("No valuation recommendation", "The workbook shows dilution sensitivity only. It does not recommend a valuation cap."),
        ("No revenue forecast", "Reservation-pricing scenarios are qualification signals, not forecast revenue or valuation support."),
        ("No tape-out budget", "Pre-seed funds ASIC feasibility and quote-backed review, not production silicon."),
        ("No customer traction claim", "Do not claim signed customers, LOIs, or revenue unless signed evidence exists."),
        ("No outcome savings claim", "Battery, heat, cooling, water, footprint, and power-bill outcomes require workload-specific measurement."),
        ("CFO/counsel review", "SAFE terms, cap, discount, pro-rata rights, option pool, tax/accounting treatment, and closing mechanics require review."),
    ]
    for i,h in enumerate(["Topic","Caveat"],1): hdr(ws8,2,i,h,fg=LIGHT)
    for r,row in enumerate(caveats,3):
        bg = ROW_A if r % 2 else ROW_B
        for c,v in enumerate(row,1): val(ws8,r,c,v,bg=bg)

    for wsx in wb.worksheets:
        for row in wsx.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal=cell.alignment.horizontal or "left", vertical="top", wrap_text=True)
        wsx.page_setup.orientation = "landscape"
        wsx.page_setup.fitToWidth = 1
        wsx.page_setup.fitToHeight = 0

    path = ROOT / "02_financial" / "ATOMiK_Financial_Model.xlsx"
    wb.save(path)
    print(f"✓ {path.name}")


def gen_financial_due_diligence_doc():
    """Convert the finance diligence memo markdown into a controlled DOCX."""
    src = ROOT / "FINANCIAL_DUE_DILIGENCE.md"
    if not src.exists():
        return
    doc = styled_doc("Financial Due Diligence Memo")
    # Remove the auto blank paragraph after title only by continuing naturally.
    lines = src.read_text().splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            # The styled title already covers the document name.
            i += 1
            continue
        if line.startswith("## "):
            h1(doc, line[3:])
            i += 1
            continue
        if line.startswith("### "):
            heading = line[4:]
            if heading == "Runway Arithmetic":
                doc.add_page_break()
            h2(doc, heading)
            i += 1
            continue
        if line.lstrip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                normalized = [c.replace(":", "").replace("-", "").strip() for c in cells]
                if all(not c for c in normalized):
                    continue
                rows.append(cells)
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        t.rows[ri].cells[ci].text = cell
                        for run in t.rows[ri].cells[ci].paragraphs[0].runs:
                            run.font.size = Pt(8.5)
                            if ri == 0:
                                run.bold = True
                doc.add_paragraph()
            continue
        if line.startswith("- "):
            bullet(doc, line[2:])
            i += 1
            continue
        if line.startswith("> "):
            body(doc, line[2:], bold=True)
            i += 1
            continue
        body(doc, line)
        i += 1
    path = ROOT / "02_financial" / "ATOMiK_Financial_Due_Diligence_Memo.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── LEGAL & FORMATION DOC ───────────────────────────────────────────────────

def gen_legal_summary():
    doc = styled_doc("ATOMiK — Legal & Formation Summary")
    body(doc, "CONFIDENTIAL — DATA ROOM DOCUMENT — NOT FOR PUBLIC DISTRIBUTION", bold=True)
    body(doc, "Last updated: May 2026. Requires counsel review before investor distribution.", size=9)
    doc.add_paragraph()

    h1(doc, "Current Entity Status")
    body(doc, "Current operating entity per internal data-room record: Rockwell Industries, LLC, California LLC, active. EIN/address details are intentionally kept out of public send materials.")
    body(doc, "Intended venture-financing path: Delaware C-Corp conversion or formation path before institutional close, pending counsel.", bold=True)
    body(doc, "Investor-safe wording: current entity and IP records exist, but final VC-ready corporate structure, cap table, and close documents require counsel review.")

    h1(doc, "Intellectual Property")
    body(doc, "Provisional patent filed — see data room 03_intellectual_property for filing details.")
    body(doc, "Patent conversion funded in pre-seed budget ($750K IP/legal allocation).")
    body(doc, "Current founder IP assignment template assigns to Rockwell Industries, LLC. Counsel must confirm whether assignment remains with the LLC, moves to a Delaware C-Corp, or is handled through conversion/assignment chain before close.")
    body(doc, "Do not state that final post-close IP assignment is complete until counsel confirms it.", bold=True)

    h1(doc, "Capitalization")
    body(doc, "See ATOMiK_Financial_Model.xlsx → Cap Table Draft sheet.")
    body(doc, "Formal cap table requires incorporation/conversion records, founder stock issuance or equivalent conversion records, option-pool decision, and counsel-reviewed financing documents.")
    body(doc, "Pre-seed instrument currently planned: SAFE, with final structure pending CFO/counsel. Valuation cap, discount, pro-rata rights, side letters, and close mechanics require CFO/counsel approval.")

    h1(doc, "Employment & Contractors")
    body(doc, "Current status for investor materials: founder-led, with first hires/fractional roles planned after financing. Any contractor/advisor arrangement should be documented before being presented as part of the company team or cap table.")
    body(doc, "Recommended before close: founder employment/consulting status reviewed by counsel, consulting agreements for any contractors, and IP assignment for all contributors.")

    h1(doc, "Before-Close Checklist")
    items = [
        ("Delaware C-Corp conversion or formation plan", "PENDING COUNSEL"),
        ("Founder stock issuance / conversion-equivalent records", "PENDING COUNSEL"),
        ("Founder IP assignment into correct post-close entity", "PENDING COUNSEL"),
        ("SAFE template and side letters", "PENDING CFO/COUNSEL"),
        ("Cap table setup and option-pool decision", "PENDING CFO/COUNSEL"),
        ("Prior obligations / advisor promises / contractor rights check", "PENDING COUNSEL"),
        ("Provisional patent filing", "FILED / SEE IP DATA ROOM"),
        ("Patent conversion plan", "FUNDED IN PRE-SEED PLAN"),
    ]
    table_2col(doc, items, header=["Document / Gate", "Status"], widths=(3.4, 3.4))

    h1(doc, "Legal / IP Attachment Index")
    table_ncol(doc, [
        ("California LLC entity record", "Available internally / controlled", "Attach or refresh before serious diligence"),
        ("Founder IP assignment template", "Available as template / counsel review pending", "Do not claim final assignment chain complete"),
        ("Provisional patent filing receipt", "Available / controlled", "Counsel-controlled or NDA path"),
        ("Provisional patent cover sheet or filing confirmation", "Available / controlled", "Use investor-safe patent summary for broad follow-up"),
        ("Patent conversion calendar", "Pending", "Track conversion deadline and counsel owner"),
        ("Cap table draft", "Pending counsel/CFO", "Do not circulate as final cap table"),
        ("SAFE template", "Pending counsel", "Do not quote finalized cap or side-letter terms"),
        ("Contractor/advisor IP assignment agreements", "None / pending / available as applicable", "Verify before close"),
        ("Counsel engagement status", "Pending / controlled note", "Update once counsel is retained for close"),
    ], header=["Attachment", "Status", "Diligence note"], font_size=8)

    path = ROOT / "03_data_room" / "legal" / "ATOMiK_Legal_Formation_Summary.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── PRODUCT / TECH DOC ──────────────────────────────────────────────────────

def gen_tech_doc():
    doc = styled_doc("ATOMiK — Product & Technical Overview")
    body(doc, "DATA ROOM — TECHNICAL SUMMARY", bold=True, size=9)

    h1(doc, "Architecture Overview")
    body(doc, "ATOMiK is a state-aware compute primitive implemented as FPGA-validated hardware IP for specific proof artifacts, with an ASIC feasibility path still to be reviewed. The core equation: state = reference_state XOR accumulated_delta")
    body(doc, "Plain-English diagram: traditional path = full state -> scan/copy/sync/rebuild -> full state again. ATOMiK path = reference state + accumulated deltas -> coalesce where possible -> READ only when needed -> SWAP boundary.", size=9)
    body(doc, "Operations:")
    table_2col(doc, [
        ("LOAD(addr, initial_state)", "Write reference state to address slot"),
        ("ACCUM(delta)", "XOR a delta into the current accumulator"),
        ("READ()", "Reconstruct current state: reference XOR accumulated deltas"),
        ("SWAP(addr)", "Commit clean epoch boundary"),
    ], widths=(2.5, 4.3))

    h1(doc, "Hardware Implementation")
    table_2col(doc, [
        ("Platform", "Zynq XC7Z020-2CLG484I (HamGeek RK-ZYNQ7020-F)"),
        ("Soft CPU", "VexRiscv SMP rv32ima in frozen Linux userspace baseline; NaxRiscv RV64GC SD-boot work remains build/bring-up context until promoted by run artifacts"),
        ("ATOMiK core", "Single-bank, 256×64-bit state table, LiteX Migen CSR"),
        ("CSR base", "0xf0000000 in frozen VexRiscv baseline; verify generated CSR map for each newer NaxRiscv/SD-boot build"),
        ("Linux", "Buildroot/Linux 6.9 in frozen userspace validation; Ubuntu/NaxRiscv work remains bring-up context unless linked to run artifacts"),
        ("Bitstream", "hamgeek_rk7020f.bit in frozen baseline; newer SD-boot bitstream should be cited by exact build artifact"),
        ("FPGA utilization", "Build-specific; do not quote utilization publicly without the matching synthesis artifact"),
    ], widths=(2.0, 4.8))

    h1(doc, "Validated Proof")
    body(doc, "Each proof item below is intentionally separated to avoid treating different evidence labels as one blended claim.", size=9)
    h2(doc, "Linux Userspace Algebraic Checks")
    body(doc, "Evidence label: HARDWARE_VALIDATED. Current summary: 16/16 algebraic property checks passed on XC7Z020 through Linux userspace.")
    h2(doc, "ATOMiK Desk v0.39-K UI")
    body(doc, "Evidence label: HARDWARE_VALIDATED. Current summary: framebuffer-native prototype on live hardware.")
    h2(doc, "AX7020 Board-Run Matrix")
    body(doc, "Evidence label: LIVE_MEASURED. Current summary: four-way comparison with raw artifacts and workload-specific caveats.")
    h2(doc, "Formal Proof Foundation")
    body(doc, "Evidence label: FORMAL_PROOF where directly audited; otherwise SOFTWARE_VALIDATED / proof work present. Do not publish theorem-count hype until reconciled.")
    h2(doc, "SD Boot Artifacts")
    body(doc, "Evidence label: BUILD_ARTIFACT. Current summary: BOOT.bin and bitstream exist; standalone power-on artifact remains gated until reproducible logs exist.")

    h1(doc, "MMIO Interface")
    body(doc, "In the frozen Linux userspace validation, ATOMiK was accessible via /dev/mem at CSR base 0xf0000000. Ordering requirement: after CSR writes, issue fence iorw,iorw plus a dummy STATUS read before reading STATE_LO/HI to flush the Wishbone pipeline. Verify the generated CSR map for newer builds.")

    h1(doc, "ASIC Path")
    body(doc, "The architecture is intended to be ASIC-portable, but ASIC feasibility remains a funded diligence task. The 256x64-bit state table in the frozen baseline is 16,384 bits, or 2KB of raw state storage; production area, power, memory macro, verification, and integration economics require expert review.")
    body(doc, "Pre-seed allocates $750K for ASIC feasibility (mentor-reviewed go/no-go). Tape-out is NOT in this round.")

    h1(doc, "Where ATOMiK Is Less Likely To Fit")
    bullet(doc, "The workload is not state-heavy.")
    bullet(doc, "Every region changes uniformly or the full state truly must move every time.")
    bullet(doc, "Hardware access overhead dominates the measured path.")
    bullet(doc, "The current software path is already optimal for the metric that matters.")
    bullet(doc, "Correctness cannot be cleanly validated.")
    bullet(doc, "The customer cannot provide a trace, baseline, or success metric.")

    h1(doc, "Diligence Questions We Expect")
    bullet(doc, "What is the exact board, build, bitstream, and commit for each artifact?")
    bullet(doc, "What is frozen validation versus current SD-boot/NaxRiscv bring-up?")
    bullet(doc, "Which proof is live measured, hardware-validated, build artifact, formal, projected, or roadmap?")
    bullet(doc, "What does the AX7020 matrix prove and what does it not prove?")
    bullet(doc, "What must happen before ASIC feasibility can be taken seriously?")
    bullet(doc, "Which customer workload is the first proof gate?")

    h1(doc, "Product Roadmap (High-Level)")
    table_2col(doc, [
        ("Now", "Zynq FPGA prototype proof paths exist. Desktop UI screenshot is hardware-validated. Workload validation is the next evidence gate."),
        ("Q3 2026", "P0 workloads targeted for hardware measurement; first paid evaluation targeted; SD boot promoted only when run artifacts exist."),
        ("Q4 2026", "IP diligence packet. Design partner agreement(s). ASIC feasibility study."),
        ("2027", "IP licensing or strategic partnership. ASIC feasibility outcome determines path."),
    ], widths=(1.5, 5.3))

    h1(doc, "Source Code & Artifacts")
    body(doc, "Private GitHub repository. Access granted under NDA.")
    bullet(doc, "hardware/ — ATOMiK RTL, Zynq integration, synthesis/build artifacts")
    bullet(doc, "hardware/zynq/ — Zynq SoC, FSBL, NaxRiscv integration")
    bullet(doc, "software/ — SDK/runtime examples and tests; cite exact test counts only after rerun")
    bullet(doc, "math/proofs/ — Lean4 formal proofs of algebraic properties")
    bullet(doc, "results/ — Raw measurement artifacts with evidence labels")

    path = ROOT / "03_data_room" / "product_tech" / "ATOMiK_Product_Technical_Overview.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── CUSTOMER PIPELINE ───────────────────────────────────────────────────────

def gen_customer_pipeline():
    doc = styled_doc("ATOMiK — Customer & Sales Pipeline")
    body(doc, "CONFIDENTIAL DATA ROOM DOCUMENT", bold=True, size=9)

    h1(doc, "Current Status")
    body(doc, "ATOMiK is pre-revenue. No signed customer contracts or LOIs exist as of May 2026. This document describes the pipeline strategy and evaluation offer structure.", bold=True)

    h1(doc, "Market Opportunity")
    table_2col(doc, [
        ("TAM context", "$1T+ 2026 semiconductor revenue backdrop. Context only; ATOMiK does not address all semiconductor spend."),
        ("SAM context", "$112B-$169B embedded systems market range from 2024 estimate to 2030 forecast; edge AI adds adjacent pressure."),
        ("Entry wedge", "$10M-$40M early annual revenue path from paid evaluations, design partners, and licensing if proof converts; not a ceiling or forecast."),
        ("IP/licensing context", "$8.14B 2025 to $11.2B 2029 semiconductor IP market context."),
    ], widths=(1.65, 5.15))

    h1(doc, "Evaluation Offer")
    body(doc, "Standard entry point: Bring one state-heavy workload, your current baseline, and the constraint that already hurts.")
    body(doc, "ATOMiK delivers: workload map, baseline comparison, evidence map, fit/no-fit recommendation, and next-step plan.")
    body(doc, "This can be structured as: (a) unpaid feasibility conversation, (b) paid technical evaluation (SOW-based), or (c) design-partner engagement with IP access.")

    h1(doc, "Target Segment Priority")
    table_2col(doc, [
        ("Edge / embedded OEMs", "Battery, heat, bandwidth-constrained products (IoT, industrial, automotive edge)"),
        ("Defense / remote systems", "Size, weight, power (SWaP) constrained; packet budget; field runtime"),
        ("AI at the edge", "Context movement, memory pressure, local inference latency"),
        ("Data center / infrastructure", "Strategic expansion path — requires edge proof first"),
    ], header=["Segment", "Pain & Entry Point"], widths=(2.2, 4.6))

    h1(doc, "Active Outreach")
    body(doc, "Investor, advisor, and design-partner outreach notes should be kept in the controlled data room. Do not present conversations as pipeline unless there is a confirmed meeting, LOI, paid evaluation, or customer artifact.")

    h1(doc, "Pipeline Tracker")
    body(doc, "Use these controlled diligence tables for updates. Keep unknowns as TBD rather than implying traction.", size=9)
    h2(doc, "Segment Strategy")
    table_ncol(doc, [
        ("Edge / embedded OEM", "Battery, heat, bandwidth", "State-heavy device workload", "Target"),
        ("Robotics / industrial edge", "Latency, reliability, control updates", "Sync/control path", "Target"),
        ("Defense-adjacent remote systems", "SWaP, packet budget, runtime", "Remote state-transfer constraint", "Target"),
        ("AI-at-the-edge context state", "Context movement, memory pressure", "Representative context-state update", "Target"),
    ], header=["Segment", "Primary pain", "Entry point", "Status"], font_size=8.0)
    h2(doc, "Controlled Tracker")
    body(doc, "Owner for current tracker rows: Founder. Current evidence for all rows: none yet unless updated with a confirmed meeting, workload, LOI, paid evaluation, or customer artifact.", size=9)
    h2(doc, "Edge / Embedded OEM")
    body(doc, "Source/status: intro, outbound, or advisor target. Next action: identify workload owner and painful constraint. Proof gate: workload map, baseline, and signed SOW. Note: do not call pipeline until meeting or artifact exists.")
    h2(doc, "Robotics / Industrial Edge")
    body(doc, "Source/status: advisor or outbound target. Next action: find a state-heavy sync/control path. Proof gate: trace, baseline, and success metric. Note: likely design-partner fit if the constraint is expensive.")
    h2(doc, "Defense-Adjacent Remote Systems")
    body(doc, "Source/status: connector or advisor target. Next action: qualify packet, radio, or SWaP constraint. Proof gate: NDA, workload map, and baseline. Note: treat as controlled; no public customer claims.")
    h2(doc, "AI-at-the-Edge Context State")
    body(doc, "Source/status: outbound or technical advisor target. Next action: frame as context-state update, not AI inference. Proof gate: representative trace and correctness oracle. Note: avoid generalized AI performance claims.")

    h1(doc, "LOI & Contract Templates")
    body(doc, "Paid evaluation proposal template: business/design_partners/paid_evaluation_proposal.md")
    body(doc, "LOI outline: business/design_partners/loi_outline.md")
    body(doc, "No signed LOIs or customer contracts exist as of this writing. First signed agreement is a key pre-seed milestone.")

    h1(doc, "Pipeline Gate")
    body(doc, "The $1.0M customer proof budget funds the first 1-3 evaluation efforts if qualified workloads are available. Success metric: measured improvement on one agreed metric while preserving correctness, with raw artifacts and a claims registry entry.")

    path = ROOT / "03_data_room" / "customers" / "ATOMiK_Customer_Pipeline.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── TEAM DOC ────────────────────────────────────────────────────────────────

def gen_team_doc():
    doc = styled_doc("ATOMiK — Team Overview")
    h1(doc, "Founder")
    body(doc, "Matthew H. Rockwell — Founder & CEO", bold=True, size=14)
    body(doc, "Email: matthew.h.rockwell@gmail.com")
    body(doc, "Role: Architecture, hardware development, software, business development. All ATOMiK proof artifacts and hardware implementation.")
    body(doc, "Background: FPGA/ASIC hardware engineering, embedded systems, state-machine architecture. Built and validated ATOMiK on Zynq FPGA hardware. Full-stack from RTL to Linux userspace.")
    body(doc, "See founder_profile.md in data room for extended profile.", size=9)

    h1(doc, "Planned First Hires (Post-Close)")
    table_2col(doc, [
        ("Senior Hardware Engineer", "ASIC-capable; RTL optimization, place-and-route, customer workload harness"),
        ("Customer Engineer / Solutions", "Technical evaluation lead; customer workload proof execution"),
        ("Fractional CFO", "SAFE terms, investor reporting, financial model, board prep"),
    ], header=["Role", "Focus"], widths=(2.5, 4.3))
    body(doc, "See first_hires.md in data room for detailed requirements.", size=9)

    h1(doc, "Advisory Board (In Progress)")
    body(doc, "Target advisors: semiconductor IP expert, ASIC/tape-out mentor, enterprise sales advisor with hardware/IP background.")
    body(doc, "No advisors formally signed as of May 2026. See advisory_board_plan.md for target criteria.", size=9)

    doc.add_page_break()
    h1(doc, "Key-Person Risk Mitigation")
    body(doc, "Senior hardware engineer — Planned post-close to support RTL optimization, customer workload harnesses, and ASIC-readiness work.")
    body(doc, "Customer / solutions engineer — Planned to execute workload proof, baseline mapping, and customer evaluation artifacts.")
    body(doc, "Fractional CFO — Needed for SAFE terms, investor reporting, financial model, and close mechanics.")
    body(doc, "IP counsel — Needed for patent conversion, assignment chain, prior-obligation review, and financing documents.")
    body(doc, "ASIC mentor / advisor — Needed before feasibility decisions and any future tape-out path.")

    h1(doc, "Gaps Before Raise Close")
    bullet(doc, "Fractional CFO for SAFE terms and financial model")
    bullet(doc, "IP counsel for patent conversion and SAFE review")
    bullet(doc, "ASIC feasibility mentor (funded in pre-seed)")

    path = ROOT / "03_data_room" / "team" / "ATOMiK_Team_Overview.docx"
    doc.save(path)
    print(f"✓ {path.name}")


# ─── COPY FILES ──────────────────────────────────────────────────────────────

def copy_files():
    # Pitch deck PPTX
    src_deck = ROOT.parent / "pitch_deck" / "ATOMiK_Investor_Deck.pptx"
    dst_deck = ROOT / "01_pitch_materials" / "ATOMiK_Investor_Deck.pptx"
    if src_deck.exists() and not dst_deck.exists():
        shutil.copy(src_deck, dst_deck)
        print(f"✓ ATOMiK_Investor_Deck.pptx (copied)")
    elif dst_deck.exists():
        print("• ATOMiK_Investor_Deck.pptx preserved")

    # Provisional patent PDF
    src_patent = ROOT.parent / "data_room" / "03_intellectual_property" / "Provisional Patent 0.0.1.pdf"
    if src_patent.exists():
        shutil.copy(src_patent, ROOT / "03_data_room" / "product_tech" / "ATOMiK_Provisional_Patent.pdf")
        print(f"✓ ATOMiK_Provisional_Patent.pdf (copied)")

    # Talking points
    src_tp = ROOT.parent / "pitch_deck" / "INVESTOR_TALKING_POINTS_FRIDAY.md"
    dst_tp = ROOT / "01_pitch_materials" / "ATOMiK_Talking_Points.md"
    if src_tp.exists() and not dst_tp.exists():
        shutil.copy(src_tp, dst_tp)
        print(f"✓ ATOMiK_Talking_Points.md (copied)")
    elif dst_tp.exists():
        print("• ATOMiK_Talking_Points.md preserved")

    # Key source markdown files for Codex reference
    src_docs = [
        ("pitch_deck/slides.md", "04_source_markdown/slides_source.md"),
        ("data_room/01_financial/financial_model.md", "04_source_markdown/financial_model_source.md"),
        ("data_room/01_financial/revenue_model_revised.md", "04_source_markdown/revenue_model_source.md"),
        ("faq/investor_faq.md", "04_source_markdown/investor_faq_source.md"),
        ("one_pager/atomik_one_pager.md", "04_source_markdown/one_pager_source.md"),
    ]
    for src_rel, dst_rel in src_docs:
        src = ROOT.parent / src_rel
        dst = ROOT / dst_rel
        if src.exists() and not dst.exists():
            shutil.copy(src, dst)
            print(f"✓ {dst_rel} (copied)")
        elif dst.exists():
            print(f"• {dst_rel} preserved")


# ─── CODEX README ────────────────────────────────────────────────────────────

def gen_codex_readme():
    readme = """# ATOMiK Investor Package — Codex Review Folder

## Purpose
This folder contains all investor-facing documents for the Friday Aggie Angel pitch.
Codex: review each document and make improvements listed below.

## What to Review

### 01_pitch_materials/
- **ATOMiK_Investor_Deck.pptx** — 16-slide deck. Review slides for clarity, flow, VC-friendly language.
  - Ensure every slide gets to the point in <5 seconds of reading
  - Remove any tech jargon a non-engineer VC wouldn't immediately understand
  - Strengthen the ROI/return narrative on Market Opportunity, Financial Model, and The Ask
  - Verify all numbers match the financial model
- **ATOMiK_Executive_Summary.docx** — 1-2 page overview. Tighten to 1 page if possible.
- **ATOMiK_Business_Plan.docx** — Full business plan. Review for accuracy and VC-friendly tone.
- **ATOMiK_One_Pager.docx** — Single page overview. Must fit on one page. Trim aggressively.
- **ATOMiK_Talking_Points.md** — Pitch script. Review for natural speech, not corporate speak.

### 02_financial/
- **ATOMiK_Financial_Due_Diligence_Memo.docx** — Founder-prepared financial diligence memo pending CFO/counsel review: safe Friday language, runway math, dilution sensitivity, revenue caveats, and remaining human-review gates.
- **ATOMiK_Financial_Model.xlsx** — Review all 12 sheets:
  - README Model Map: intended use, assumptions, investor-facing/internal sheet map
  - Use of Funds: formula-backed totals and proof-gated categories
  - Runway & Burn: monthly budget and reserve sensitivity
  - 18-Month Cash Plan: monthly target-plan burn by spend line
  - Evaluation Pricing: reservation prices as qualification signals, not forecast revenue
  - Evaluation SOW Economics: planning-only SOW economics, not forecast revenue
  - Milestone Gates: 18-month proof gates tied to investor evidence
  - Dilution Sensitivity: illustrative SAFE cap math, not a cap recommendation
  - Market Context: verify all source citations are accurate
  - Cap Table Draft: clearly mark as DRAFT and avoid false precision
  - Unmodeled CFO Needs: taxes, benefits, insurance, legal close, tooling, travel, contingency
  - Caveats: ensure CFO/counsel review boundaries are explicit
- Flag any numbers that need CFO approval before sharing

### 03_data_room/
- **legal/ATOMiK_Legal_Formation_Summary.docx** — Ensure all PENDING items are clearly flagged
- **product_tech/ATOMiK_Product_Technical_Overview.docx** — Make accessible to non-engineers
- **customers/ATOMiK_Customer_Pipeline.docx** — Honest status of pipeline
- **team/ATOMiK_Team_Overview.docx** — Professional, factual

## Codex Instructions

1. **DO NOT** change any financial numbers without flagging them for human review
2. **DO NOT** remove evidence labels (HARDWARE_VALIDATED, LIVE_MEASURED, etc.)
3. **DO NOT** add performance claims that aren't in the claims_registry.yaml
4. **DO** improve prose clarity, grammar, flow, and VC-appropriate language
5. **DO** ensure consistent formatting within each document
6. **DO** strengthen the investor ROI narrative where appropriate
7. **DO** flag any inconsistencies between documents
8. **DO** note any PENDING items that need human action before Friday

## Key Claims Rules (DO NOT VIOLATE)
- Only claim battery/heat/water/cooling savings if there's a measured artifact
- Always pair performance numbers with evidence labels
- Never say "production ready" or "commercial product"
- Always say what IS measured vs what is an evaluation target

## Evidence Labels Currently Active
- HARDWARE_VALIDATED: v0.39-K UI, algebraic tests, Linux userspace path
- LIVE_MEASURED: AX7020 board run matrix
- FORMAL_PROOF: directly audited formal algebra properties
- SOFTWARE_VALIDATED: proof work present where not independently audited
- BUILD_ARTIFACT: SD boot artifacts

## Source Files
- 04_source_markdown/ contains the canonical markdown sources
- Website source is in /website/src/
- Claims registry is in /results/claims_registry.yaml

## Contact
matthew.h.rockwell@gmail.com
"""
    with open(ROOT / "CODEX_README.md", "w") as f:
        f.write(readme)
    print("✓ CODEX_README.md")


# ─── MAIN ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating ATOMiK investor package...\n")
    gen_executive_summary()
    gen_business_plan()
    gen_one_pager()
    gen_financial_model()
    gen_financial_due_diligence_doc()
    gen_legal_summary()
    gen_tech_doc()
    gen_customer_pipeline()
    gen_team_doc()
    copy_files()
    gen_codex_readme()
    print("\n✓ All documents generated in business/for_codex/")
    print("  Open ATOMiK_Investor_Deck.pptx to verify visual quality.")
    print("  Open ATOMiK_Financial_Model.xlsx to verify numbers.")
    print("  All .docx files ready for Codex review.")
